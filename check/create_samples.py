# create_samples.py
import os
from docx import Document

os.makedirs("data/good_docs", exist_ok=True)
os.makedirs("data/poor_docs", exist_ok=True)

# 잘 쓴 파일 내용
good_text = {
    "② 처분의 내용": """귀 공단 ○○지사로부터 2025년 10월 4일 자 건강보험료 부과 처분을 받았습니다.
해당 처분은 2023년도 귀속 종합소득 신고 내역을 기준으로 산정된 것으로 안내받았으나,
실제 소득자료 중 일부가 반영되지 않아 과도한 보험료가 부과된 것으로 판단됩니다.
이에 따라 정확한 소득금액 확인 및 보험료 재산정을 요청드립니다.""",

    "④ 이의신청의 취지와 사유": """본인은 위 부과 처분이 실제 소득과 불일치하여 과다 산정된 것으로 보입니다.
종합소득세 신고서 및 원천징수영수증을 첨부하오니,
관련 자료를 근거로 보험료 산정 과정을 재검토하여 주시기 바랍니다.
아울러 해당 기간의 소득 정정 결과에 따라 적정 보험료로 조정될 수 있도록 요청드립니다."""
}

# 못 쓴 파일 내용
poor_text = {
    "② 처분의 내용": """건강보험료가 너무 많이 나왔습니다.
공단에서 잘못 계산한 것 같아요.
작년에 소득이 줄었는데 그대로 나와서 이상합니다.""",

    "④ 이의신청의 취지와 사유": """돈이 너무 많이 나왔고, 계산이 잘못된 것 같습니다.
다시 계산해 주셨으면 합니다.
증빙자료는 나중에 제출하겠습니다."""
}

# 잘 쓴 파일 작성
good_doc = Document()
good_doc.add_heading("② 처분의 내용", level=2)
good_doc.add_paragraph(good_text["② 처분의 내용"])
good_doc.add_heading("④ 이의신청의 취지와 사유", level=2)
good_doc.add_paragraph(good_text["④ 이의신청의 취지와 사유"])
good_doc.save("data/good_docs/good_example.docx")

# 못 쓴 파일 작성
poor_doc = Document()
poor_doc.add_heading("② 처분의 내용", level=2)
poor_doc.add_paragraph(poor_text["② 처분의 내용"])
poor_doc.add_heading("④ 이의신청의 취지와 사유", level=2)
poor_doc.add_paragraph(poor_text["④ 이의신청의 취지와 사유"])
poor_doc.save("data/poor_docs/poor_example.docx")

print("✅ 샘플 문서가 성공적으로 생성되었습니다.")
print(" - data/good_docs/good_example.docx")
print(" - data/poor_docs/poor_example.docx")
